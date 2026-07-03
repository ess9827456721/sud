"""Flask application factory for Трекер судебных дел v2."""
import logging
import sqlite3
import threading
from datetime import datetime, timedelta

from flask import Flask, g, jsonify, redirect, render_template, request, flash, url_for

from court_tracker.config import DB_PATH, FLASK_HOST, FLASK_PORT, FLASK_DEBUG
from court_tracker.db.schema import init_db, run_migrations
from court_tracker.db import queries


logger = logging.getLogger(__name__)


def _save_events_and_deadlines(conn, case_id: int, events: list) -> None:
    """Save events then auto-create deadline entries for the case."""
    queries.save_events(conn, case_id, events)
    try:
        from court_tracker.services.deadline_service import auto_create_deadlines_for_case
        auto_create_deadlines_for_case(conn, case_id)
    except Exception as exc:
        logger.warning("auto_create_deadlines error case %s: %s", case_id, exc)

# ---------------------------------------------------------------------------
# Sync state (shared across threads)
# ---------------------------------------------------------------------------
_sync_lock = threading.Lock()
_sync_state = {"is_running": False, "last_sync": None}


def _background_sync(app: Flask) -> None:
    with app.app_context():
        _sync_state["is_running"] = True
        try:
            conn = _get_db()
            cases = queries.get_all_cases(conn, {"source": "kad"})
            from court_tracker.scraper.kad_scraper import KADScraper
            for case in cases:
                if not case.get("kad_url"):
                    continue
                try:
                    with KADScraper() as scraper:
                        details = scraper.get_case_details(case["kad_url"])
                    if details:
                        details["case_number"] = case["case_number"]
                        cid = queries.upsert_case(conn, details)
                        queries.save_participants(conn, cid, details.get("participants", []))
                        _save_events_and_deadlines(conn, cid, details.get("events", []))
                        queries.log_sync(conn, cid, True, "startup sync")
                except Exception as exc:
                    queries.log_sync(conn, case["id"], False, str(exc))
            _sync_state["last_sync"] = datetime.utcnow().isoformat()
        except Exception as exc:
            logger.error("background sync error: %s", exc)
        finally:
            _sync_state["is_running"] = False


# ---------------------------------------------------------------------------
# DB helpers
# ---------------------------------------------------------------------------

def _get_db() -> sqlite3.Connection:
    if "db" not in g:
        conn = sqlite3.connect(str(DB_PATH))
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA foreign_keys=ON")
        conn.row_factory = sqlite3.Row
        init_db(conn)
        run_migrations(conn)
        g.db = conn
    return g.db


def _close_db(exc=None) -> None:
    conn = g.pop("db", None)
    if conn:
        conn.close()


# ---------------------------------------------------------------------------
# Application factory
# ---------------------------------------------------------------------------

def create_app() -> Flask:
    app = Flask(__name__, template_folder="templates", static_folder="static")
    app.secret_key = "court-tracker-dev-secret"

    app.teardown_appcontext(_close_db)

    # Scheduler (Phase 7)
    from court_tracker.scraper.scheduler import SyncScheduler
    scheduler = SyncScheduler()
    app.extensions["scheduler"] = scheduler

    # Startup sync
    @app.before_request
    def _ensure_db():
        _get_db()

    with app.app_context():
        conn = _get_db()
        interval_h = float(queries.get_setting(conn, "sync_interval_hours") or 2)
        scheduler.update_interval(interval_h)
        if queries.get_setting(conn, "sync_on_startup") == "1":
            scheduler.start(delay_seconds=5)

    # ------------------------------------------------------------------
    # Template helpers
    # ------------------------------------------------------------------

    @app.template_filter("fmt_money")
    def fmt_money(v) -> str:
        if v is None:
            return "—"
        try:
            f = float(v)
            if f == 0:
                return "—"
            return f"{f:,.0f} ₽".replace(",", " ")
        except (TypeError, ValueError):
            return "—"

    @app.template_filter("urgency_class")
    def urgency_class(deadline_date_str: str) -> str:
        try:
            d = datetime.strptime(deadline_date_str[:10], "%Y-%m-%d").date()
        except (ValueError, TypeError):
            return "urgency-safe"
        days = (d - datetime.utcnow().date()).days
        if days <= 1:
            return "urgency-crit"
        if days <= 3:
            return "urgency-soon"
        if days <= 7:
            return "urgency-warn"
        return "urgency-safe"

    @app.template_filter("days_until")
    def days_until(date_str: str) -> int:
        try:
            d = datetime.strptime(date_str[:10], "%Y-%m-%d").date()
            return (d - datetime.utcnow().date()).days
        except (ValueError, TypeError):
            return 999

    @app.context_processor
    def inject_globals():
        conn = _get_db()
        unread = conn.execute(
            "SELECT COUNT(*) FROM notifications WHERE is_read=0"
        ).fetchone()[0]
        return {"unread_count": unread, "now": datetime.utcnow()}

    # ------------------------------------------------------------------
    # Dashboard
    # ------------------------------------------------------------------

    @app.route("/")
    def index():
        conn = _get_db()
        all_cases = queries.get_all_cases(conn)
        this_month = datetime.utcnow().strftime("%Y-%m")
        active = [c for c in all_cases if c["status"] not in ("Завершено", "Прекращено", None)]
        new_this_month = [c for c in all_cases if (c.get("created_at") or "")[:7] == this_month]
        kad_count = sum(1 for c in all_cases if c["source"] == "kad")
        soy_count = sum(1 for c in all_cases if c["source"] == "soy")

        upcoming_hearings = queries.get_upcoming_hearings(conn, days=14)
        approaching_deadlines = queries.get_approaching_deadlines(conn, days=14)
        expiring_poa = queries.get_expiring_powers_of_attorney(conn, days=30)

        recent_logs = conn.execute(
            """SELECT sl.*, c.case_number FROM sync_log sl
               LEFT JOIN cases c ON c.id = sl.case_id
               ORDER BY sl.synced_at DESC LIMIT 5"""
        ).fetchall()

        kanban_counts = dict(
            conn.execute(
                "SELECT stage, COUNT(*) FROM kanban_stage GROUP BY stage"
            ).fetchall()
        )

        fin = queries.get_dashboard_financials(conn)
        failed_soy = queries.get_failed_soy_cases(conn)
        return render_template(
            "index.html",
            total_active=len(active),
            new_this_month=len(new_this_month),
            kad_count=kad_count,
            soy_count=soy_count,
            upcoming_hearings=upcoming_hearings,
            approaching_deadlines=approaching_deadlines,
            expiring_poa=expiring_poa,
            recent_logs=recent_logs,
            kanban_counts=kanban_counts,
            total_active_claims=fin["total_active_claims"],
            fee_receivable=fin["fee_receivable"],
            failed_soy=failed_soy,
        )

    # ------------------------------------------------------------------
    # Cases
    # ------------------------------------------------------------------

    @app.route("/cases")
    def cases_list():
        conn = _get_db()
        filters = {
            k: request.args.get(k)
            for k in ("source", "status", "client_id")
            if request.args.get(k)
        }
        case_list = queries.get_all_cases(conn, filters or None)
        clients = conn.execute("SELECT id, name FROM clients ORDER BY name").fetchall()
        statuses = conn.execute(
            "SELECT DISTINCT status FROM cases WHERE status IS NOT NULL ORDER BY status"
        ).fetchall()
        return render_template(
            "cases.html",
            cases=case_list,
            clients=clients,
            statuses=statuses,
            filters=filters,
        )

    @app.route("/cases/<int:case_id>")
    def case_detail(case_id: int):
        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if not case:
            flash("Дело не найдено.", "danger")
            return redirect(url_for("cases_list"))
        clients = conn.execute("SELECT id, name FROM clients ORDER BY name").fetchall()
        return render_template("case_detail.html", case=case, clients=clients)

    @app.route("/cases/add", methods=["GET", "POST"])
    def case_add():
        conn = _get_db()
        clients = conn.execute("SELECT id, name FROM clients ORDER BY name").fetchall()
        if request.method == "POST":
            tab = request.form.get("tab", "kad")
            if tab == "kad":
                case_number = request.form.get("case_number", "").strip()
                if not case_number:
                    flash("Введите номер дела.", "danger")
                    return render_template("add_case.html", clients=clients, active_tab="kad")
                from court_tracker.scraper.utils import normalize_case_number
                from court_tracker.scraper.kad_scraper import KADScraper
                case_number = normalize_case_number(case_number)
                import sqlite3 as _sqlite3
                try:
                    from playwright.sync_api import TimeoutError as PWTimeoutError
                except Exception:  # playwright not installed
                    class PWTimeoutError(Exception):
                        pass
                try:
                    with KADScraper() as scraper:
                        result = scraper.search_by_case_number(case_number)
                        details = None
                        if result and result.get("kad_url"):
                            details = scraper.get_case_details(result["kad_url"])
                    if not result:
                        flash(f"Дело не найдено: {case_number}", "error")
                        return render_template(
                            "add_case.html", clients=clients, active_tab="kad",
                            scraper_error="not found: " + case_number,
                        )
                    data = details if details else result
                    data.setdefault("case_number", case_number)
                    data.setdefault("source", "kad")
                    case_id = queries.upsert_case(conn, data)
                    if details:
                        queries.save_participants(conn, case_id, details.get("participants", []))
                        _save_events_and_deadlines(conn, case_id, details.get("events", []))
                    queries.log_sync(conn, case_id, True, "manual add-case")
                    flash(f"Дело {case_number} добавлено.", "success")
                    return redirect(url_for("case_detail", case_id=case_id))
                except PWTimeoutError as exc:
                    flash("Сайт КАД не ответил вовремя. Попробуйте позже.", "error")
                    return render_template("add_case.html", clients=clients, active_tab="kad",
                                           scraper_error="Timeout: " + str(exc)[:200])
                except _sqlite3.IntegrityError as exc:
                    flash("Ошибка базы данных при сохранении дела.", "error")
                    return render_template("add_case.html", clients=clients, active_tab="kad",
                                           scraper_error="UNIQUE constraint: " + str(exc)[:200])
                except Exception as exc:
                    flash(f"Ошибка парсинга: {exc}", "error")
                    return render_template("add_case.html", clients=clients, active_tab="kad",
                                           scraper_error=str(exc)[:200])
            else:
                # SOY manual entry
                f = request.form
                data = {
                    "case_number": f.get("case_number", "").strip(),
                    "source": "soy",
                    "court": f.get("court"),
                    "judge": f.get("judge"),
                    "status": f.get("status"),
                    "case_type": f.get("case_type"),
                    "start_date": f.get("start_date"),
                    "soy_url_first": f.get("soy_url_first"),
                    "soy_url_appeal": f.get("soy_url_appeal"),
                    "soy_url_cassation": f.get("soy_url_cassation"),
                    "custom_label": f.get("custom_label"),
                    "client_id": f.get("client_id") or None,
                }
                if not data["case_number"]:
                    flash("Введите номер дела.", "danger")
                    return render_template("add_case.html", clients=clients, active_tab="soy")
                case_id = queries.upsert_case(conn, data)
                # Add plaintiff / respondent as participants
                for role, name_key in (("Истец", "plaintiff_name"), ("Ответчик", "respondent_name")):
                    name = f.get(name_key, "").strip()
                    if name:
                        exists = conn.execute(
                            "SELECT 1 FROM participants WHERE case_id=? AND role=? AND name=?",
                            (case_id, role, name),
                        ).fetchone()
                        if not exists:
                            conn.execute(
                                "INSERT INTO participants(case_id, role, name) VALUES (?,?,?)",
                                (case_id, role, name),
                            )
                conn.commit()
                flash(f"Дело {data['case_number']} добавлено (СОЮ).", "success")
                return redirect(url_for("case_detail", case_id=case_id))
        return render_template("add_case.html", clients=clients, active_tab="kad")

    @app.route("/cases/<int:case_id>/delete", methods=["POST"])
    def case_delete(case_id: int):
        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if case:
            conn.execute("DELETE FROM cases WHERE id = ?", (case_id,))
            conn.commit()
            flash(f"Дело {case['case_number']} удалено.", "success")
        return redirect(url_for("index"))

    @app.route("/cases/<int:case_id>/sync", methods=["POST"])
    def case_sync(case_id: int):
        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if not case or not case.get("kad_url"):
            flash("Нет URL для синхронизации.", "warning")
            return redirect(url_for("case_detail", case_id=case_id))
        import sqlite3 as _sqlite3
        try:
            from playwright.sync_api import TimeoutError as PWTimeoutError
        except Exception:
            class PWTimeoutError(Exception):
                pass
        scraper_error = None
        try:
            from court_tracker.scraper.kad_scraper import KADScraper
            with KADScraper() as scraper:
                details = scraper.get_case_details(case["kad_url"])
            if details:
                details["case_number"] = case["case_number"]
                queries.upsert_case(conn, details)
                queries.save_participants(conn, case_id, details.get("participants", []))
                _save_events_and_deadlines(conn, case_id, details.get("events", []))
                queries.log_sync(conn, case_id, True, "manual sync")
                flash("Синхронизация выполнена.", "success")
            else:
                queries.log_sync(conn, case_id, False, "scraper returned None")
                flash("Ошибка синхронизации.", "error")
                scraper_error = "scraper: не удалось получить данные дела"
        except PWTimeoutError as exc:
            queries.log_sync(conn, case_id, False, "Timeout: " + str(exc)[:200])
            flash("Сайт КАД не ответил вовремя.", "error")
            scraper_error = "Timeout: " + str(exc)[:200]
        except _sqlite3.IntegrityError as exc:
            queries.log_sync(conn, case_id, False, "IntegrityError: " + str(exc)[:200])
            flash("Ошибка базы данных.", "error")
            scraper_error = "UNIQUE constraint: " + str(exc)[:200]
        except Exception as exc:
            queries.log_sync(conn, case_id, False, str(exc)[:200])
            flash(f"Ошибка: {exc}", "error")
            scraper_error = str(exc)[:200]
        if scraper_error:
            case = queries.get_case_full(conn, case_id)
            clients = conn.execute("SELECT id, name FROM clients ORDER BY name").fetchall()
            return render_template("case_detail.html", case=case, clients=clients,
                                   scraper_error=scraper_error)
        return redirect(url_for("case_detail", case_id=case_id))

    # ------------------------------------------------------------------
    # Settings
    # ------------------------------------------------------------------

    @app.route("/settings", methods=["GET", "POST"])
    def settings():
        conn = _get_db()
        sched = app.extensions.get("scheduler")
        if request.method == "POST":
            new_interval = request.form.get("sync_interval_hours", "2")
            queries.set_setting(conn, "sync_interval_hours", new_interval)
            queries.set_setting(conn, "sync_on_startup", "1" if request.form.get("sync_on_startup") else "0")
            if sched:
                sched.update_interval(float(new_interval))
            flash("Настройки сохранены.", "success")
            return redirect(url_for("settings"))
        current = {
            "sync_interval_hours": queries.get_setting(conn, "sync_interval_hours", "2"),
            "sync_on_startup": queries.get_setting(conn, "sync_on_startup", "1"),
            "app_version": queries.get_setting(conn, "app_version", "2.0"),
        }
        sync_status = sched.get_status() if sched else {}
        failed_soy = queries.get_failed_soy_cases(conn)
        return render_template("settings.html", settings=current,
                               sync_status=sync_status, failed_soy=failed_soy)

    # ------------------------------------------------------------------
    # JSON API
    # ------------------------------------------------------------------

    @app.route("/api/autosave", methods=["POST"])
    def api_autosave():
        data = request.get_json(force=True, silent=True) or {}
        table = data.get("table", "")
        rec_id = data.get("id")
        field = data.get("field", "")
        value = data.get("value")
        if not table or rec_id is None or not field:
            return jsonify({"success": False, "error": "missing fields"}), 400
        from court_tracker.services.autosave import autosave_field
        ok = autosave_field(table, int(rec_id), field, value)
        if ok:
            return jsonify({"success": True})
        return jsonify({"success": False, "error": "field not allowed or record not found"}), 400

    @app.route("/api/logs/last")
    def logs_last():
        """Return last 100 sync_log records as plain text (used by error overlay)."""
        conn = _get_db()
        rows = conn.execute(
            """SELECT sl.synced_at, sl.success, sl.message, c.case_number
               FROM sync_log sl
               LEFT JOIN cases c ON c.id = sl.case_id
               ORDER BY sl.id DESC LIMIT 100"""
        ).fetchall()
        lines = []
        for r in rows:
            status = "OK" if r["success"] else "ERR"
            case = r["case_number"] or "(общее)"
            lines.append(f"[{r['synced_at']}] {status} {case}: {r['message']}")
        return "\n".join(lines), 200, {"Content-Type": "text/plain; charset=utf-8"}

    @app.route("/api/cases")
    def api_cases():
        conn = _get_db()
        return jsonify(queries.get_all_cases(conn))

    @app.route("/api/cases/<int:case_id>")
    def api_case(case_id: int):
        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if not case:
            return jsonify({"error": "not found"}), 404
        return jsonify(dict(case))

    @app.route("/api/upcoming")
    def api_upcoming():
        conn = _get_db()
        return jsonify(queries.get_upcoming_hearings(conn, days=30))

    @app.route("/api/deadlines")
    def api_deadlines():
        conn = _get_db()
        return jsonify(queries.get_approaching_deadlines(conn, days=14))

    @app.route("/api/sync/<int:case_id>", methods=["POST"])
    def api_sync(case_id: int):
        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if not case or not case.get("kad_url"):
            return jsonify({"success": False, "message": "no KAD URL"}), 400
        try:
            from court_tracker.scraper.kad_scraper import KADScraper
            with KADScraper() as scraper:
                details = scraper.get_case_details(case["kad_url"])
            if details:
                details["case_number"] = case["case_number"]
                queries.upsert_case(conn, details)
                queries.save_participants(conn, case_id, details.get("participants", []))
                _save_events_and_deadlines(conn, case_id, details.get("events", []))
                queries.log_sync(conn, case_id, True, "api sync")
                return jsonify({"success": True, "message": "synced"})
            queries.log_sync(conn, case_id, False, "scraper None")
            return jsonify({"success": False, "message": "scraper returned no data"})
        except Exception as exc:
            queries.log_sync(conn, case_id, False, str(exc))
            return jsonify({"success": False, "message": str(exc)}), 500

    @app.route("/api/sync/status/legacy")
    def api_sync_status_legacy():
        conn = _get_db()
        return jsonify({
            "last_sync": _sync_state["last_sync"],
            "is_running": _sync_state["is_running"],
            "cases_count": conn.execute("SELECT COUNT(*) FROM cases").fetchone()[0],
        })

    @app.route("/api/notifications")
    def api_notifications():
        conn = _get_db()
        rows = conn.execute(
            "SELECT * FROM notifications WHERE is_read=0 ORDER BY created_at DESC LIMIT 50"
        ).fetchall()
        return jsonify([dict(r) for r in rows])

    @app.route("/api/notifications/read-all", methods=["POST"])
    def api_notifications_read_all():
        conn = _get_db()
        conn.execute("UPDATE notifications SET is_read=1")
        conn.commit()
        return jsonify({"success": True})

    # ------------------------------------------------------------------
    # Clients — HTML routes
    # ------------------------------------------------------------------

    @app.route("/clients")
    def clients_list():
        conn = _get_db()
        clients = queries.get_all_clients(conn)
        return render_template("clients.html", clients=clients)

    @app.route("/clients/add", methods=["GET", "POST"])
    def client_add():
        if request.method == "POST":
            conn = _get_db()
            f = request.form
            data = {
                "type":           f.get("type", "legal"),
                "name":           f.get("name", "").strip(),
                "short_name":     f.get("short_name", "").strip() or None,
                "inn":            f.get("inn", "").strip() or None,
                "ogrn":           f.get("ogrn", "").strip() or None,
                "address":        f.get("address", "").strip() or None,
                "status_egrul":   f.get("status_egrul") or None,
                "phone":          f.get("phone", "").strip() or None,
                "email":          f.get("email", "").strip() or None,
                "contact_person": f.get("contact_person", "").strip() or None,
                "notes":          f.get("notes", "").strip() or None,
            }
            if not data["name"]:
                flash("Укажите наименование клиента.", "danger")
                return render_template("client_form.html", client=None, action="add")
            client_id = queries.upsert_client(conn, data)
            flash("Клиент добавлен.", "success")
            return redirect(url_for("client_detail", client_id=client_id))
        return render_template("client_form.html", client=None, action="add")

    @app.route("/clients/<int:client_id>")
    def client_detail(client_id: int):
        conn = _get_db()
        client = queries.get_client_with_cases(conn, client_id)
        if not client:
            flash("Клиент не найден.", "danger")
            return redirect(url_for("clients_list"))
        all_cases = queries.get_all_cases(conn)
        return render_template("client_detail.html", client=client, all_cases=all_cases)

    @app.route("/clients/<int:client_id>/edit", methods=["GET", "POST"])
    def client_edit(client_id: int):
        conn = _get_db()
        client = queries.get_client_with_cases(conn, client_id)
        if not client:
            flash("Клиент не найден.", "danger")
            return redirect(url_for("clients_list"))
        if request.method == "POST":
            f = request.form
            data = {
                "type":           f.get("type", client["type"]),
                "name":           f.get("name", "").strip() or client["name"],
                "short_name":     f.get("short_name", "").strip() or None,
                "inn":            client["inn"],  # INN is the upsert key — don't change
                "ogrn":           f.get("ogrn", "").strip() or None,
                "address":        f.get("address", "").strip() or None,
                "status_egrul":   f.get("status_egrul") or None,
                "phone":          f.get("phone", "").strip() or None,
                "email":          f.get("email", "").strip() or None,
                "contact_person": f.get("contact_person", "").strip() or None,
                "notes":          f.get("notes", "").strip() or None,
            }
            # upsert_client matches by INN; fall back to direct UPDATE when INN is missing
            if data["inn"]:
                queries.upsert_client(conn, data)
            else:
                now = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%S")
                conn.execute(
                    """UPDATE clients SET type=?,name=?,short_name=?,ogrn=?,address=?,
                       status_egrul=?,phone=?,email=?,contact_person=?,notes=?,updated_at=?
                       WHERE id=?""",
                    (data["type"], data["name"], data["short_name"], data["ogrn"],
                     data["address"], data["status_egrul"], data["phone"], data["email"],
                     data["contact_person"], data["notes"], now, client_id),
                )
                conn.commit()
            flash("Клиент обновлён.", "success")
            return redirect(url_for("client_detail", client_id=client_id))
        return render_template("client_form.html", client=client, action="edit")

    @app.route("/clients/<int:client_id>/delete", methods=["POST"])
    def client_delete(client_id: int):
        conn = _get_db()
        active = conn.execute(
            "SELECT COUNT(*) FROM cases WHERE client_id=? AND status NOT IN ('Завершено','Прекращено')",
            (client_id,),
        ).fetchone()[0]
        if active > 0:
            flash(f"Нельзя удалить клиента: есть {active} активных дел.", "danger")
            return redirect(url_for("client_detail", client_id=client_id))
        conn.execute("DELETE FROM clients WHERE id=?", (client_id,))
        conn.commit()
        flash("Клиент удалён.", "success")
        return redirect(url_for("clients_list"))

    # ------------------------------------------------------------------
    # Client contacts
    # ------------------------------------------------------------------

    @app.route("/clients/<int:client_id>/contacts/add", methods=["POST"])
    def contact_add(client_id: int):
        conn = _get_db()
        f = request.form
        queries.add_contact(conn, {
            "client_id": client_id,
            "name":  f.get("name", "").strip(),
            "role":  f.get("role", "").strip(),
            "phone": f.get("phone", "").strip(),
            "email": f.get("email", "").strip(),
            "notes": f.get("notes", "").strip(),
        })
        flash("Контакт добавлен.", "success")
        return redirect(url_for("client_detail", client_id=client_id) + "#tab-contacts")

    @app.route("/clients/<int:client_id>/contacts/<int:contact_id>/delete", methods=["POST"])
    def contact_delete(client_id: int, contact_id: int):
        conn = _get_db()
        queries.delete_contact(conn, contact_id)
        flash("Контакт удалён.", "success")
        return redirect(url_for("client_detail", client_id=client_id) + "#tab-contacts")

    # ------------------------------------------------------------------
    # Powers of attorney
    # ------------------------------------------------------------------

    @app.route("/clients/<int:client_id>/poa/add", methods=["POST"])
    def poa_add(client_id: int):
        conn = _get_db()
        f = request.form
        file = request.files.get("scan")
        attachment_id = None
        if file and file.filename:
            import os, uuid
            from court_tracker.config import ATTACHMENTS_DIR
            ext = os.path.splitext(file.filename)[1]
            stored = f"poa_{uuid.uuid4().hex}{ext}"
            dest = ATTACHMENTS_DIR / stored
            file.save(str(dest))
            # We need a case_id for attachments FK — use NULL pattern via dummy case
            # PoA scans are stored with case_id=0 sentinel (FK allows NULL, use 0 workaround)
            # Actually the schema has case_id NOT NULL; store against the first linked case or skip
            # For now: find any case linked to this client
            row = conn.execute("SELECT id FROM cases WHERE client_id=? LIMIT 1", (client_id,)).fetchone()
            if row:
                cur = conn.execute(
                    """INSERT INTO attachments(case_id,filename,stored_name,file_path,file_size,mime_type)
                       VALUES (?,?,?,?,?,?)""",
                    (row[0], file.filename, stored, str(dest),
                     os.path.getsize(str(dest)), file.content_type),
                )
                attachment_id = cur.lastrowid
                conn.commit()
        queries.add_power_of_attorney(conn, {
            "client_id":        client_id,
            "case_id":          f.get("case_id") or None,
            "number":           f.get("number", "").strip() or None,
            "issue_date":       f.get("issue_date") or None,
            "expiry_date":      f.get("expiry_date") or None,
            "scope_description":f.get("scope_description", "").strip() or None,
            "attachment_id":    attachment_id,
        })
        flash("Доверенность добавлена.", "success")
        return redirect(url_for("client_detail", client_id=client_id) + "#tab-poa")

    @app.route("/clients/<int:client_id>/poa/<int:poa_id>/delete", methods=["POST"])
    def poa_delete(client_id: int, poa_id: int):
        conn = _get_db()
        queries.delete_power_of_attorney(conn, poa_id)
        flash("Доверенность удалена.", "success")
        return redirect(url_for("client_detail", client_id=client_id) + "#tab-poa")

    # ------------------------------------------------------------------
    # Client JSON API
    # ------------------------------------------------------------------

    @app.route("/api/clients")
    def api_clients():
        conn = _get_db()
        return jsonify(queries.get_all_clients(conn))

    @app.route("/api/clients", methods=["POST"])
    def api_client_create():
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        if not data.get("name"):
            return jsonify({"error": "name required"}), 400
        client_id = queries.upsert_client(conn, data)
        client = queries.get_client_with_cases(conn, client_id)
        return jsonify(dict(client)), 201

    @app.route("/api/egrul/<inn>")
    def api_egrul(inn: str):
        from court_tracker.services.egrul_service import fetch_by_inn
        data = fetch_by_inn(inn)
        if data is None:
            return jsonify({"error": "not found"}), 404
        return jsonify(data)

    @app.route("/api/cases/<int:case_id>/set-client", methods=["POST"])
    def api_set_client(case_id: int):
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        client_id = data.get("client_id")
        now = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%S")
        conn.execute(
            "UPDATE cases SET client_id=?, updated_at=? WHERE id=?",
            (client_id, now, case_id),
        )
        conn.commit()
        return jsonify({"success": True})

    # ------------------------------------------------------------------
    # Deadlines API (Phase 4)
    # ------------------------------------------------------------------

    @app.route("/api/cases/<int:case_id>/deadlines")
    def api_case_deadlines(case_id: int):
        from court_tracker.services.deadline_service import urgency_level, DEADLINE_RULES
        conn = _get_db()
        deadlines = queries.get_deadlines_for_case(conn, case_id)
        for d in deadlines:
            d["urgency"] = urgency_level(d.get("deadline_date"))
        return jsonify(deadlines)

    @app.route("/api/cases/<int:case_id>/deadlines", methods=["POST"])
    def api_create_deadline(case_id: int):
        from court_tracker.services.deadline_service import (
            calculate_deadline, DEADLINE_RULES
        )
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}

        # If a rule_key + trigger_date are provided, auto-calculate
        rule_key = data.get("rule_key")
        trigger_date_str = data.get("trigger_date")
        if rule_key and trigger_date_str and rule_key in DEADLINE_RULES:
            from datetime import date as _date
            try:
                td = _date.fromisoformat(trigger_date_str[:10])
            except ValueError:
                return jsonify({"error": "invalid trigger_date"}), 400
            dl_date = calculate_deadline(td, rule_key)
            rule = DEADLINE_RULES[rule_key]
            data.setdefault("deadline_type", rule["name"])
            data.setdefault("statute_reference", rule["statute_reference"])
            data.setdefault("statute_article", rule["statute_article"])
            data.setdefault("calculation_note", rule["calculation_note"])
            data.setdefault("deadline_date", dl_date.isoformat() if dl_date else None)
            data["is_auto_calculated"] = 1
            data["trigger_event"] = rule["trigger_event"]

        data["case_id"] = case_id
        if not data.get("deadline_date"):
            return jsonify({"error": "deadline_date required"}), 400

        dl_id = queries.create_deadline(conn, data)
        return jsonify({"success": True, "id": dl_id, "deadline_date": data["deadline_date"]}), 201

    @app.route("/api/deadlines/<int:dl_id>", methods=["PATCH"])
    def api_update_deadline(dl_id: int):
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        new_date = data.get("deadline_date")
        if not new_date:
            return jsonify({"error": "deadline_date required"}), 400
        ok = queries.update_deadline_date(conn, dl_id, new_date)
        if not ok:
            return jsonify({"error": "not found"}), 404
        return jsonify({"success": True, "new_date": new_date, "calculation_note": "изменено вручную"})

    @app.route("/api/deadlines/<int:dl_id>", methods=["DELETE"])
    def api_delete_deadline(dl_id: int):
        conn = _get_db()
        ok = queries.delete_deadline(conn, dl_id)
        if not ok:
            return jsonify({"error": "not found"}), 404
        return jsonify({"success": True})

    @app.route("/api/deadlines/<int:dl_id>/done", methods=["PATCH"])
    def api_toggle_deadline_done(dl_id: int):
        conn = _get_db()
        new_state = queries.toggle_deadline_done(conn, dl_id)
        if new_state is None:
            return jsonify({"error": "not found"}), 404
        return jsonify({"success": True, "is_done": new_state})

    # Helper: expose DEADLINE_RULES to front-end
    @app.route("/api/deadline-rules")
    def api_deadline_rules():
        from court_tracker.services.deadline_service import DEADLINE_RULES
        return jsonify(DEADLINE_RULES)

    # ------------------------------------------------------------------
    # Kanban (Phase 4)
    # ------------------------------------------------------------------

    @app.route("/kanban")
    def kanban():
        conn = _get_db()
        board = queries.get_kanban_board(conn)
        # Attach critical deadline dot info
        for stage, cases in board.items():
            for c in cases:
                c["critical"] = bool(queries.get_case_critical_deadline(conn, c["id"]))
        return render_template(
            "kanban.html",
            board=board,
            stage_names=queries.STAGE_NAMES,
            stages=queries.KANBAN_STAGES,
        )

    @app.route("/api/kanban")
    def api_kanban():
        conn = _get_db()
        board = queries.get_kanban_board(conn)
        return jsonify(board)

    @app.route("/api/kanban/move", methods=["POST"])
    def api_kanban_move():
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        case_id = data.get("case_id")
        new_stage = data.get("new_stage")
        if not case_id or not new_stage:
            return jsonify({"error": "case_id and new_stage required"}), 400
        if new_stage not in queries.KANBAN_STAGES:
            return jsonify({"error": f"invalid stage: {new_stage}"}), 400
        warning = queries.move_kanban(conn, int(case_id), new_stage)
        resp = {"success": True}
        if warning:
            resp["mismatch_warning"] = warning
        return jsonify(resp)

    # ------------------------------------------------------------------
    # Notes & Tasks (Phase 5)
    # ------------------------------------------------------------------

    @app.route("/api/cases/<int:case_id>/notes")
    def api_get_notes(case_id: int):
        conn = _get_db()
        return jsonify(queries.get_notes_for_case(conn, case_id))

    @app.route("/api/cases/<int:case_id>/notes", methods=["POST"])
    def api_create_note(case_id: int):
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        data["case_id"] = case_id
        note_id = queries.create_note(conn, data)
        all_notes = queries.get_notes_for_case(conn, case_id)
        new_note = next((n for n in all_notes if n["id"] == note_id),
                        {"id": note_id, "item_type": data.get("item_type", "note"),
                         "checklist": [], "tags": []})
        return jsonify(new_note), 201

    @app.route("/api/notes/<int:note_id>", methods=["PUT"])
    def api_update_note(note_id: int):
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        queries.update_note(conn, note_id, data)
        return jsonify({"success": True})

    @app.route("/api/notes/<int:note_id>", methods=["DELETE"])
    def api_delete_note(note_id: int):
        conn = _get_db()
        ok = queries.delete_note(conn, note_id)
        if not ok:
            return jsonify({"error": "not found"}), 404
        return jsonify({"success": True})

    @app.route("/api/notes/<int:note_id>/checklist", methods=["POST"])
    def api_add_checklist(note_id: int):
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        text = (data.get("text") or "").strip()
        if not text:
            return jsonify({"error": "text required"}), 400
        item_id = queries.add_checklist_item(conn, note_id, text)
        return jsonify({"success": True, "id": item_id, "text": text,
                        "checked": False, "position": 0}), 201

    @app.route("/api/checklist/<int:item_id>/toggle", methods=["PATCH"])
    def api_toggle_checklist(item_id: int):
        conn = _get_db()
        new_state = queries.toggle_checklist_item(conn, item_id)
        if new_state is None:
            return jsonify({"error": "not found"}), 404
        return jsonify({"success": True, "checked": new_state})

    @app.route("/api/checklist/<int:item_id>", methods=["DELETE"])
    def api_delete_checklist(item_id: int):
        conn = _get_db()
        ok = queries.delete_checklist_item(conn, item_id)
        if not ok:
            return jsonify({"error": "not found"}), 404
        return jsonify({"success": True})

    @app.route("/api/notes/<int:note_id>/tags", methods=["POST"])
    def api_add_tag(note_id: int):
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        tag = (data.get("tag") or "").strip()
        if not tag:
            return jsonify({"error": "tag required"}), 400
        queries.add_note_tag(conn, note_id, tag)
        return jsonify({"success": True, "tag": tag})

    @app.route("/api/notes/<int:note_id>/tags", methods=["DELETE"])
    def api_delete_tag(note_id: int):
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        tag = (data.get("tag") or "").strip()
        if not tag:
            return jsonify({"error": "tag required"}), 400
        queries.delete_note_tag(conn, note_id, tag)
        return jsonify({"success": True})

    @app.route("/api/tags")
    def api_tags():
        conn = _get_db()
        return jsonify(queries.get_all_tags(conn))

    @app.route("/api/notes/reorder", methods=["POST"])
    def api_reorder_notes():
        conn = _get_db()
        data = request.get_json(force=True, silent=True) or {}
        case_id = data.get("case_id")
        ordered_ids = data.get("ordered_ids", [])
        if not case_id or not ordered_ids:
            return jsonify({"error": "case_id and ordered_ids required"}), 400
        queries.reorder_notes(conn, int(case_id), [int(i) for i in ordered_ids])
        return jsonify({"success": True})

    # ------------------------------------------------------------------
    # Attachments (Phase 5)
    # ------------------------------------------------------------------

    from court_tracker.config import ALLOWED_EXTENSIONS, MAX_FILE_SIZE_MB

    def _allowed_file(filename: str) -> bool:
        return "." in filename and \
               filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

    @app.route("/api/cases/<int:case_id>/attachments")
    def api_get_attachments(case_id: int):
        conn = _get_db()
        return jsonify(queries.get_attachments_for_case(conn, case_id))

    @app.route("/api/cases/<int:case_id>/attachments", methods=["POST"])
    def api_upload_attachment(case_id: int):
        import os
        import uuid
        from flask import send_file as _sf
        from court_tracker.config import ATTACHMENTS_DIR

        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if not case:
            return jsonify({"error": "case not found"}), 404

        if "file" not in request.files:
            return jsonify({"error": "no file"}), 400
        file = request.files["file"]
        if not file.filename:
            return jsonify({"error": "empty filename"}), 400
        if not _allowed_file(file.filename):
            return jsonify({"error": "file type not allowed"}), 400

        file.seek(0, 2)
        size = file.tell()
        file.seek(0)
        if size > MAX_FILE_SIZE_MB * 1024 * 1024:
            return jsonify({"error": f"file too large (max {MAX_FILE_SIZE_MB} MB)"}), 400

        ext = os.path.splitext(file.filename)[1].lower()
        stored = f"{uuid.uuid4().hex}{ext}"
        dest = ATTACHMENTS_DIR / stored
        file.save(str(dest))

        note_id_raw = request.form.get("note_id")
        note_id = int(note_id_raw) if note_id_raw else None

        att_id = queries.create_attachment(conn, {
            "case_id":     case_id,
            "note_id":     note_id,
            "filename":    file.filename,
            "stored_name": stored,
            "file_path":   str(dest),
            "file_size":   size,
            "mime_type":   file.content_type,
        })
        all_att = queries.get_attachments_for_case(conn, case_id)
        att = next((a for a in all_att if a["id"] == att_id), {"id": att_id})
        return jsonify(att), 201

    @app.route("/api/cases/<int:case_id>/attachments/<int:att_id>/download")
    def api_download_attachment(case_id: int, att_id: int):
        from flask import send_file as _sf
        conn = _get_db()
        row = conn.execute(
            "SELECT * FROM attachments WHERE id=? AND case_id=?", (att_id, case_id)
        ).fetchone()
        if not row:
            return jsonify({"error": "not found"}), 404
        att = dict(row)
        return _sf(att["file_path"], as_attachment=True,
                   download_name=att["filename"])

    @app.route("/api/attachments/<int:att_id>", methods=["DELETE"])
    def api_delete_attachment(att_id: int):
        import os
        conn = _get_db()
        file_path = queries.delete_attachment(conn, att_id)
        if not file_path:
            return jsonify({"error": "not found"}), 404
        try:
            os.remove(file_path)
        except OSError:
            pass
        return jsonify({"success": True})

    # ------------------------------------------------------------------
    # Export (Phase 7)
    # ------------------------------------------------------------------

    @app.route("/export/cases.xlsx")
    def export_cases_xlsx():
        import io
        from flask import send_file as _sf
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        conn = _get_db()
        cases    = queries.get_all_cases_for_export(conn)
        future_ev = queries.get_future_events_for_export(conn)
        open_dl  = queries.get_open_deadlines_for_export(conn)

        wb = openpyxl.Workbook()

        # ── Sheet 1: Cases ────────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Реестр дел"
        hdr_fill = PatternFill("solid", fgColor="2B4EFF")
        hdr_font = Font(color="FFFFFF", bold=True)
        headers1 = [
            "ID", "Номер дела", "Суд", "Судья", "Статус", "Тип дела", "Источник",
            "Начало", "Клиент", "Сумма иска", "Взыскано", "Итог",
            "URL КАД", "Обновлено",
        ]
        for col, h in enumerate(headers1, 1):
            cell = ws1.cell(row=1, column=col, value=h)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal="center")
        for r, c in enumerate(cases, 2):
            ws1.append([
                c["id"], c["case_number"], c["court"], c["judge"], c["status"],
                c["case_type"], c["source"], c["start_date"],
                c.get("client_name"), c.get("claim_amount"), c.get("awarded_amount"),
                c.get("outcome"), c.get("kad_url"), (c["updated_at"] or "")[:16],
            ])
        for col in ws1.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            ws1.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        # ── Sheet 2: Future events ─────────────────────────────────────────
        ws2 = wb.create_sheet("Заседания")
        headers2 = ["Дата", "Тип", "Суд", "Номер дела", "Описание"]
        for col, h in enumerate(headers2, 1):
            cell = ws2.cell(row=1, column=col, value=h)
            cell.font = Font(bold=True)
        for ev in future_ev:
            ws2.append([
                ev["event_date"], ev["event_type"],
                ev.get("court"), ev["case_number"],
                ev.get("description"),
            ])
        for col in ws2.columns:
            ws2.column_dimensions[col[0].column_letter].width = 30

        # ── Sheet 3: Open deadlines ────────────────────────────────────────
        ws3 = wb.create_sheet("Сроки")
        headers3 = ["Срок до", "Тип срока", "Номер дела", "Суд", "Статья", "Примечание"]
        for col, h in enumerate(headers3, 1):
            cell = ws3.cell(row=1, column=col, value=h)
            cell.font = Font(bold=True)
        for dl in open_dl:
            ws3.append([
                dl["deadline_date"], dl["deadline_type"],
                dl["case_number"], dl.get("court"),
                dl.get("statute_article"), dl.get("calculation_note"),
            ])
        for col in ws3.columns:
            ws3.column_dimensions[col[0].column_letter].width = 30

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return _sf(buf, as_attachment=True,
                   download_name="cases_export.xlsx",
                   mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    @app.route("/export/cases/<int:case_id>/report.docx")
    def export_case_report(case_id: int):
        import io
        from flask import send_file as _sf
        from docx import Document as _Doc
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if not case:
            return "Дело не найдено", 404

        client_name = ""
        if case.get("client_id"):
            row = conn.execute("SELECT name FROM clients WHERE id=?", (case["client_id"],)).fetchone()
            if row:
                client_name = row[0]

        notes = queries.get_notes_for_report(conn, case_id)

        doc = _Doc()
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # Title
        title = doc.add_heading(f"Отчёт по делу {case.get('case_number', '')}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Section helper
        def _section(name):
            h = doc.add_heading(name, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x2B, 0x4E, 0xFF)

        def _table_row(table, label, value):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value or "—")
            row.cells[0].paragraphs[0].runs[0].bold = True

        # Section 1: General info
        _section("Общие сведения")
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for label, val in [
            ("Суд",           case.get("court")),
            ("Судья",         case.get("judge")),
            ("Статус",        case.get("status")),
            ("Тип дела",      case.get("case_type")),
            ("Дата начала",   case.get("start_date")),
            ("Клиент",        client_name),
            ("Сумма иска",    case.get("claim_amount")),
            ("Взыскано",      case.get("awarded_amount")),
            ("Итог",          case.get("outcome")),
        ]:
            _table_row(t, label, val)
        doc.add_paragraph()

        # Section 2: Participants
        _section("Участники")
        parts = case.get("participants", [])
        if parts:
            t2 = doc.add_table(rows=1, cols=4)
            t2.style = "Table Grid"
            for i, h in enumerate(["Роль", "Наименование", "ИНН", "Адрес"]):
                t2.rows[0].cells[i].text = h
                t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for p in parts:
                row = t2.add_row()
                row.cells[0].text = p.get("role") or ""
                row.cells[1].text = p.get("name") or ""
                row.cells[2].text = p.get("inn") or ""
                row.cells[3].text = p.get("address") or ""
        else:
            doc.add_paragraph("Участники не указаны.")
        doc.add_paragraph()

        # Section 3: Events
        _section("Хронология событий")
        events = sorted(case.get("events", []), key=lambda e: e.get("event_date") or "")
        if events:
            t3 = doc.add_table(rows=1, cols=3)
            t3.style = "Table Grid"
            for i, h in enumerate(["Дата", "Тип", "Описание"]):
                t3.rows[0].cells[i].text = h
                t3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for ev in events:
                row = t3.add_row()
                row.cells[0].text = ev.get("event_date") or ""
                row.cells[1].text = ev.get("event_type") or ""
                row.cells[2].text = ev.get("description") or ""
        else:
            doc.add_paragraph("События не найдены.")
        doc.add_paragraph()

        # Section 4: Deadlines
        _section("Процессуальные сроки")
        deadlines = sorted(case.get("deadlines", []), key=lambda d: d.get("deadline_date") or "")
        if deadlines:
            t4 = doc.add_table(rows=1, cols=4)
            t4.style = "Table Grid"
            for i, h in enumerate(["Срок", "Дата", "Статья", "Выполнен"]):
                t4.rows[0].cells[i].text = h
                t4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for dl in deadlines:
                row = t4.add_row()
                row.cells[0].text = dl.get("deadline_type") or ""
                row.cells[1].text = dl.get("deadline_date") or ""
                row.cells[2].text = dl.get("statute_article") or ""
                row.cells[3].text = "Да" if dl.get("is_done") else "Нет"
        else:
            doc.add_paragraph("Сроки не указаны.")
        doc.add_paragraph()

        # Section 5: Notes
        plain_notes = [n for n in notes if n.get("item_type") != "task"]
        if plain_notes:
            _section("Заметки")
            for n in plain_notes:
                doc.add_heading(n.get("title") or "Заметка", level=2)
                if n.get("body"):
                    doc.add_paragraph(n["body"])
                for ci in n.get("checklist", []):
                    mark = "☑" if ci.get("checked") else "☐"
                    doc.add_paragraph(f"{mark} {ci.get('text', '')}", style="List Bullet")
            doc.add_paragraph()

        # Section 6: Completed tasks
        done_tasks = [n for n in notes if n.get("item_type") == "task" and n.get("task_status") == "done"]
        if done_tasks:
            _section("Выполненные мероприятия")
            for tk in done_tasks:
                doc.add_paragraph(
                    f"✔ {tk.get('title') or 'Задача'}  "
                    f"(приоритет: {tk.get('task_priority', '')}, "
                    f"дедлайн: {tk.get('task_due_date') or '—'})",
                    style="List Bullet",
                )
            doc.add_paragraph()

        # Footer
        doc.add_paragraph(
            f"Сформировано: {datetime.utcnow().strftime('%d.%m.%Y %H:%M')} UTC"
        ).alignment = WD_ALIGN_PARAGRAPH.RIGHT

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_num = (case.get("case_number") or "case").replace("/", "-")
        return _sf(buf, as_attachment=True,
                   download_name=f"report_{safe_num}.docx",
                   mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    @app.route("/export/backup.zip")
    def export_backup():
        import io, shutil, zipfile
        from flask import send_file as _sf
        from court_tracker.config import DATA_DIR

        conn = _get_db()
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            db_path = DATA_DIR / "court_tracker.db"
            if db_path.exists():
                zf.write(str(db_path), f"backup_{ts}/court_tracker.db")
            att_dir = DATA_DIR / "attachments"
            if att_dir.exists():
                for att_file in att_dir.rglob("*"):
                    if att_file.is_file():
                        zf.write(str(att_file),
                                 f"backup_{ts}/attachments/{att_file.name}")
        buf.seek(0)
        queries.log_sync(conn, None, True, f"Backup created: backup_{ts}.zip")
        return _sf(buf, as_attachment=True,
                   download_name=f"backup_{ts}.zip",
                   mimetype="application/zip")

    # ------------------------------------------------------------------
    # SOY scraping API (Phase 7)
    # ------------------------------------------------------------------

    @app.route("/api/cases/<int:case_id>/soy-sync", methods=["POST"])
    def api_soy_sync(case_id: int):
        conn = _get_db()
        case = queries.get_case_full(conn, case_id)
        if not case:
            return jsonify({"success": False, "error": "not found"}), 404

        url = (case.get("soy_url_cassation")
               or case.get("soy_url_appeal")
               or case.get("soy_url_first"))
        if not url:
            return jsonify({"success": False, "error": "no SOY URL configured"}), 400

        from court_tracker.scraper.soy_scraper import SOYScraper
        scraper = SOYScraper(headless=True)
        result = scraper.scrape_case(url)

        if result["success"]:
            _save_events_and_deadlines(conn, case_id, result["events"])
            if result.get("participants"):
                queries.save_participants(conn, case_id, result["participants"], smart=True)
            # Save extended case_info fields from SOY
            ci = result.get("case_info") or {}
            _SOY_SAFE_FIELDS = {
                "judge", "case_category", "case_category_path",
                "receipt_date", "decision_date", "decision_result",
                "court_first", "uid",
            }
            for field, val in ci.items():
                if field in _SOY_SAFE_FIELDS and val:
                    queries.upsert_case_field(conn, case_id, field, val)
            queries.update_soy_scrape_status(conn, case_id, "success")
            queries.log_sync(conn, case_id, True, "manual СОЮ sync")
            return jsonify({"success": True, "status": "success",
                            "message": "Данные успешно обновлены",
                            "events_count": len(result["events"])})
        else:
            queries.update_soy_scrape_status(conn, case_id,
                                             result["status"], result.get("error_msg"))
            queries.log_sync(conn, case_id, False, result.get("error_msg", ""))
            return jsonify({"success": False, "status": result["status"],
                            "message": result.get("error_msg", "Ошибка синхронизации")})

    @app.route("/api/cases/<int:case_id>/soy-scraping", methods=["POST"])
    def api_soy_toggle(case_id: int):
        data = request.get_json(force=True, silent=True) or {}
        enabled = bool(data.get("enabled", False))
        conn = _get_db()
        conn.execute(
            "UPDATE cases SET soy_scraping_enabled=?, updated_at=datetime('now') WHERE id=?",
            (1 if enabled else 0, case_id),
        )
        conn.commit()
        return jsonify({"success": True, "enabled": enabled})

    # ------------------------------------------------------------------
    # Participants — manual/freeze API (Phase 7.2)
    # ------------------------------------------------------------------

    @app.route("/api/participants/<int:part_id>", methods=["PATCH"])
    def api_participant_edit(part_id: int):
        """Manually edit a participant field and freeze it (_manual=1)."""
        data = request.get_json(force=True, silent=True) or {}
        field = data.get("field", "")
        value = data.get("value", "")
        if field not in ("inn", "name", "address"):
            return jsonify({"success": False, "error": "field must be inn, name or address"}), 400
        conn = _get_db()
        ok = queries.freeze_participant_field(conn, part_id, field, value)
        if ok:
            return jsonify({"success": True, "frozen": True, "field": field})
        return jsonify({"success": False, "error": "participant not found"}), 404

    @app.route("/api/participants/<int:part_id>/unfreeze", methods=["POST"])
    def api_participant_unfreeze(part_id: int):
        """Remove manual freeze from one or all fields of a participant."""
        data = request.get_json(force=True, silent=True) or {}
        field = data.get("field")  # None → unfreeze all
        conn = _get_db()
        queries.unfreeze_participant_field(conn, part_id, field)
        return jsonify({"success": True, "field": field or "all"})

    # ------------------------------------------------------------------
    # Scheduler API (Phase 7)
    # ------------------------------------------------------------------

    @app.route("/api/sync/status")
    def api_sync_status():
        sched = app.extensions.get("scheduler")
        if sched:
            return jsonify(sched.get_status())
        return jsonify({"is_running": False, "last_sync": None, "next_sync": None, "cases_count": 0})

    @app.route("/api/sync/trigger", methods=["POST"])
    def api_sync_trigger():
        sched = app.extensions.get("scheduler")
        if sched:
            sched.trigger_now()
        return jsonify({"success": True, "message": "Синхронизация запущена"})

    # ------------------------------------------------------------------
    # Analytics (Phase 6)
    # ------------------------------------------------------------------

    @app.route("/analytics")
    def analytics():
        return render_template("analytics.html")

    @app.route("/api/analytics/cases")
    def api_analytics_cases():
        conn = _get_db()
        return jsonify(queries.get_analytics_cases(conn))

    @app.route("/api/analytics/judges")
    def api_analytics_judges():
        conn = _get_db()
        return jsonify(queries.get_analytics_judges(conn))

    @app.route("/api/analytics/finance")
    def api_analytics_finance():
        conn = _get_db()
        return jsonify(queries.get_analytics_finance(conn))

    # ------------------------------------------------------------------
    # Calendar (Phase 6)
    # ------------------------------------------------------------------

    @app.route("/calendar")
    def calendar():
        return render_template("calendar.html")

    @app.route("/api/calendar")
    def api_calendar():
        conn = _get_db()
        try:
            year  = int(request.args.get("year",  datetime.utcnow().year))
            month = int(request.args.get("month", datetime.utcnow().month))
        except ValueError:
            return jsonify({"error": "invalid year/month"}), 400
        if not (1 <= month <= 12):
            return jsonify({"error": "month out of range"}), 400
        return jsonify(queries.get_calendar_events(conn, year, month))

    # ------------------------------------------------------------------
    # Document Templates (Phase 5)
    # ------------------------------------------------------------------

    @app.route("/templates")
    def templates_list():
        import os
        from court_tracker.config import TEMPLATES_DIR
        files = sorted(
            f for f in os.listdir(str(TEMPLATES_DIR))
            if f.lower().endswith(".docx")
        )
        return jsonify([{"filename": f} for f in files])

    @app.route("/templates/generate", methods=["POST"])
    def templates_generate():
        import os
        import uuid
        from flask import send_file as _sf
        from court_tracker.config import TEMPLATES_DIR, TEMP_DIR

        data = request.get_json(force=True, silent=True) or {}
        template_filename = data.get("template_filename", "")
        case_id = data.get("case_id")

        if not template_filename or not case_id:
            return jsonify({"error": "template_filename and case_id required"}), 400
        if "/" in template_filename or "\\" in template_filename \
                or not template_filename.lower().endswith(".docx"):
            return jsonify({"error": "invalid filename"}), 400

        template_path = TEMPLATES_DIR / template_filename
        if not template_path.exists():
            return jsonify({"error": "template not found"}), 404

        conn = _get_db()
        case = queries.get_case_full(conn, int(case_id))
        if not case:
            return jsonify({"error": "case not found"}), 404

        participants = case.get("participants", [])
        plaintiff  = next(
            (p["name"] for p in participants
             if p.get("role") and "истец" in p["role"].lower()), "")
        respondent = next(
            (p["name"] for p in participants
             if p.get("role") and "ответчик" in p["role"].lower()), "")
        client_name = ""
        if case.get("client_id"):
            cl_row = conn.execute(
                "SELECT name FROM clients WHERE id=?", (case["client_id"],)
            ).fetchone()
            if cl_row:
                client_name = cl_row[0]

        placeholders = {
            "{{номер_дела}}": case.get("case_number") or "",
            "{{суд}}":        case.get("court") or "",
            "{{судья}}":      case.get("judge") or "",
            "{{истец}}":      plaintiff,
            "{{ответчик}}":   respondent,
            "{{дата}}":       datetime.utcnow().strftime("%d.%m.%Y"),
            "{{клиент}}":     client_name,
            "{{сумма_иска}}": "",
        }

        def _replace_in_para(para):
            for key, val in placeholders.items():
                if key in para.text:
                    for run in para.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, val)

        try:
            from docx import Document as _DocxDocument
            doc = _DocxDocument(str(template_path))
            for para in doc.paragraphs:
                _replace_in_para(para)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_in_para(para)
            out_name = f"{uuid.uuid4().hex}_{template_filename}"
            out_path = TEMP_DIR / out_name
            doc.save(str(out_path))
        except Exception as exc:
            return jsonify({"error": f"generation failed: {exc}"}), 500

        from flask import send_file as _sf
        safe_case_number = (case.get("case_number") or "case").replace("/", "-")
        return _sf(str(out_path), as_attachment=True,
                   download_name=f"{safe_case_number}_{template_filename}")

    return app


# ---------------------------------------------------------------------------
# Run directly
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    app = create_app()
    app.run(host=FLASK_HOST, port=FLASK_PORT, debug=FLASK_DEBUG, threaded=True)
